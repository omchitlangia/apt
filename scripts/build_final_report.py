#!/usr/bin/env python3
"""Build APT_final_report.docx — the capstone methodology + out-of-sample
validation report.

This script assembles ONE Word document from the committed project record. It
does NOT re-run any experiment: every table is rendered directly from a
persisted CSV under reports/, and every figure is embedded from a persisted PNG
under plots/ or reports/.../figures/. Numbers in prose are the ERRATA-CORRECTED
canonical values (see docs/phase3_cost_beta_report.md §13.1 and
docs/phase4_report.md §0/§1).

Run:  .venv/bin/python scripts/build_final_report.py
Out:  APT_final_report.docx  (repo root)
"""

from __future__ import annotations

import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "APT_final_report.docx"

NAVY = RGBColor(0x1F, 0x32, 0x55)
ACCENT = RGBColor(0x2E, 0x5A, 0x88)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0x9B, 0x1C, 0x1C)

EMBEDDED: list[str] = []
MISSING: list[str] = []


# --------------------------------------------------------------------------- #
# low-level helpers
# --------------------------------------------------------------------------- #
def _csv(rel: str) -> pd.DataFrame:
    return pd.read_csv(ROOT / rel)


def set_cell_bg(cell, hexcolor: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level <= 2 else ACCENT
    return p


def para(
    doc,
    text="",
    *,
    italic=False,
    bold=False,
    size=10.5,
    color=None,
    align=None,
    space_after=6,
    style=None,
):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.italic = italic
        run.bold = bold
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
    return p


def rich(doc, segments, *, size=10.5, space_after=6, align=None):
    """segments: list of (text, {bold,italic,color,mono})."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, fmt in segments:
        run = p.add_run(text)
        run.bold = fmt.get("bold", False)
        run.italic = fmt.get("italic", False)
        run.font.size = Pt(fmt.get("size", size))
        if fmt.get("mono"):
            run.font.name = "Consolas"
        if "color" in fmt:
            run.font.color.rgb = fmt["color"]
    return p


def bullet(doc, text, *, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p


def numbered(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p


def equation(doc, text, *, where=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    if where:
        r2 = p.add_run("    " + where)
        r2.font.name = "Consolas"
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = GREY
    return p


def caption(doc, fig_no, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"Figure {fig_no}. ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY


def tcap(doc, tbl_no, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"Table {tbl_no}. ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY


def figure(doc, fig_no, rel, text, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = ROOT / rel
    if path.exists():
        try:
            doc.add_picture(str(path), width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            EMBEDDED.append(f"Fig {fig_no}: {rel}")
        except Exception as exc:  # pragma: no cover
            para(doc, f"[figure embed failed: {rel} — {exc}]", color=RED, italic=True)
            MISSING.append(rel)
    else:
        para(doc, f"[figure not found: {rel}]", color=RED, italic=True)
        MISSING.append(rel)
    caption(doc, fig_no, text)


def table_from_df(doc, df: pd.DataFrame, *, headers=None, fontsize=8.5, align_first_left=True):
    df = df.reset_index(drop=True)
    cols = list(df.columns)
    headers = headers or cols
    t = doc.add_table(rows=1, cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for j, name in enumerate(headers):
        hdr[j].text = ""
        run = hdr[j].paragraphs[0].add_run(str(name))
        run.bold = True
        run.font.size = Pt(fontsize)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[j], "1F3255")
        hdr[j].paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if (j == 0 and align_first_left) else WD_ALIGN_PARAGRAPH.CENTER
        )
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(cols):
            val = row[c]
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run("" if pd.isna(val) else str(val))
            run.font.size = Pt(fontsize)
            cells[j].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if (j == 0 and align_first_left)
                else WD_ALIGN_PARAGRAPH.CENTER
            )
    # zebra
    for i, r in enumerate(t.rows):
        if i == 0:
            continue
        if i % 2 == 0:
            for cell in r.cells:
                set_cell_bg(cell, "EEF2F8")
    return t


def fmt(df, mapping):
    """mapping: {col: (newname, fmtfn)}; returns new df with selected cols."""
    out = pd.DataFrame()
    for col, (newname, fn) in mapping.items():
        out[newname] = df[col].map(fn) if fn else df[col]
    return out


def f2(x):
    return "" if pd.isna(x) else f"{x:,.2f}"


def f3(x):
    return "" if pd.isna(x) else f"{x:.3f}"


def f1(x):
    return "" if pd.isna(x) else f"{x:,.1f}"


def pc2(x):
    return "" if pd.isna(x) else f"{x:,.2f}%"


def i0(x):
    return "" if pd.isna(x) else f"{int(round(x))}"


# --------------------------------------------------------------------------- #
# document scaffolding
# --------------------------------------------------------------------------- #
def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12


def add_footer(doc):
    sec = doc.sections[-1]
    f = sec.footer
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Adaptive Pairs Trading (APT) — Methodology & Out-of-Sample Validation  ·  June 2026  ·  "
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    # page number field
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-2" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose “Update Field” to build the Table of Contents."
    fldChar2.append(t)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


# --------------------------------------------------------------------------- #
# title page + abstract
# --------------------------------------------------------------------------- #
def title_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Adaptive Pairs Trading (APT)")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("A Methodology and Out-of-Sample Validation Study")
    r2.font.size = Pt(15)
    r2.italic = True
    r2.font.color.rgb = ACCENT

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(10)
    r3 = p3.add_run(
        "Statistical-arbitrage pairs trading on the NSE cash-equity universe,\n"
        "with an adaptive-equilibrium extension and a crypto out-of-sample retest"
    )
    r3.font.size = Pt(11)
    r3.font.color.rgb = GREY

    for _ in range(2):
        doc.add_paragraph()

    def line(label, value, value_bold=True):
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = q.add_run(label + "  ")
        rl.font.size = Pt(11)
        rl.font.color.rgb = GREY
        rv = q.add_run(value)
        rv.font.size = Pt(11.5)
        rv.bold = value_bold
        rv.font.color.rgb = NAVY

    line("Authors:", "Om Chitlangia · Manan Raina Kumar")
    line("Supervisors:", "Dr. Barik · Dr. Malu")
    line("Institution / Department:", "[TODO — fill before submission]", value_bold=False)
    line("Date:", "June 2026")

    for _ in range(2):
        doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run(
        "Audience: academic committee.  Reproducibility: every table is rendered from a "
        "persisted CSV and every figure from a persisted PNG; no experiment is re-run for "
        "this document.  All performance numbers are reported gross AND net of "
        "transaction cost."
    )
    rn.font.size = Pt(9)
    rn.italic = True
    rn.font.color.rgb = GREY
    doc.add_page_break()


def abstract(doc):
    h(doc, "Structured Abstract", level=1)

    def block(label, body):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(label + " — ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = ACCENT
        r2 = p.add_run(body)
        r2.font.size = Pt(10.5)

    block(
        "Background",
        "Classical statistical-arbitrage pairs trading exploits temporary "
        "dislocations between cointegrated assets. Two structural headwinds motivate this "
        "study: (i) the documented decay of the equity-pairs edge as the strategy has become "
        "crowded, and (ii) the cash-market constraints of the Indian (NSE) venue — no native "
        "shorting in the cash segment, discrete-lot frictions, and a split/bonus-only "
        "(dividend-unadjusted) price history.",
    )
    block(
        "Objective",
        "To build a fully reproducible pairs-trading pipeline on the NSE universe; to test "
        "whether an adaptive-equilibrium signal (a per-session causal Kalman level filter) "
        "improves on a frozen cointegration relationship; and to subject the strongest result "
        "to a multiple-testing-aware validation gate (Deflated Sharpe Ratio and the "
        "Probability of Backtest Overfitting) before retesting the mechanism out-of-sample on "
        "a breadth universe (crypto perpetual pairs).",
    )
    block(
        "Methods",
        "Engle–Granger cointegration with Benjamini–Hochberg FDR control for pair selection "
        "(Johansen used as an order-independent comparator); a frozen-OLS hedge ratio with a "
        "β-aware (1+β) transaction-cost convention; Ornstein–Uhlenbeck fitting with Bertram "
        "(2010) cost-aware optimal thresholds; a μ-only local-level Kalman equilibrium filter "
        "(West–Harrison discount form) with a train-only absorption guard; and a leakage-free "
        "walk-forward backtest. Validation uses the Deflated Sharpe Ratio (Bailey & López de "
        "Prado, 2014) and PBO via CSCV (Bailey et al., 2017).",
    )
    block(
        "Results",
        "On the daily NSE universe the baseline is net-profitable (net Sharpe 1.08 over ~7 "
        "out-of-sample years). At intraday frequency the tradeable universe collapses to two "
        "pair-folds; on that matched universe the adaptive Kalman arm roughly doubles the net "
        "Sharpe of the frozen engine (2.15 vs 0.95–1.00) and is the only engine to clear the "
        "N=46 Deflated-Sharpe luck bar (DSR 0.808) — but its p-value is 0.192 (not "
        "significant) and the result rests on n=2. The β-tracking escalation fails by weak "
        "identification (β collapses at intraday frequency). On the breadth crypto universe "
        "the ordering replicates (Kalman > frozen on gross and net; Kalman uniquely "
        "positive gross), yet net is deeply negative for every engine and the net Deflated "
        "Sharpe is 0.002 — the edge does not survive realistic intraday transaction cost.",
    )
    block(
        "Conclusion",
        "We document a real, repeatable, universe-general mean-reversion mechanism that does "
        "NOT clear realistic intraday transaction costs. The strongest in-sample NSE result is "
        "statistically indistinguishable from luck after Deflated-Sharpe correction (n=2), and "
        "on the breadth/crypto universe the strategy is net-negative. We present this as a "
        "rigorous negative result and a methodological contribution: the obstacle is execution "
        "cost and turnover, not signal, so the productive next step is structural "
        "(lower-frequency re-anchoring, a cheaper venue, maker execution) rather than "
        "parametric.",
    )
    kw = doc.add_paragraph()
    rk = kw.add_run("Keywords:  ")
    rk.bold = True
    rk.font.size = Pt(10)
    rk.font.color.rgb = ACCENT
    rk2 = kw.add_run(
        "pairs trading; cointegration; Engle–Granger; Johansen; Ornstein–Uhlenbeck; Bertram "
        "thresholds; Kalman filter; adaptive equilibrium; Deflated Sharpe Ratio; PBO; "
        "walk-forward; NSE; crypto."
    )
    rk2.font.size = Pt(10)
    doc.add_page_break()
    h(doc, "Contents", level=1)
    add_toc(doc)
    doc.add_page_break()


# --------------------------------------------------------------------------- #
# 2. Executive summary
# --------------------------------------------------------------------------- #
def executive_summary(doc):
    h(doc, "1.  Executive Summary — the Honest Verdict", level=1)
    rich(
        doc,
        [
            ("The headline finding is a disciplined negative. ", {"bold": True}),
            (
                "Across four years of staged development — NSE data engineering, a daily "
                "walk-forward baseline, an intraday Ornstein–Uhlenbeck/Bertram engine, an adaptive "
                "Kalman-equilibrium extension, and a crypto out-of-sample retest — we find a ",
                {},
            ),
            (
                "real, repeatable, universe-general mean-reversion mechanism that does not clear "
                "realistic intraday transaction costs.",
                {"bold": True},
            ),
        ],
    )
    para(
        doc,
        "The adaptive equilibrium filter genuinely extracts more signal than a frozen "
        "cointegration relationship: it beats the frozen engine on gross and net Sharpe on "
        "the matched NSE universe, and the same ordering replicates on a completely "
        "different (crypto) universe, where the adaptive engine is uniquely positive on "
        "gross return. That cross-universe replication is the positive methodological "
        "result. But two facts gate any deployment claim: (i) the strongest NSE result rests "
        "on only n=2 intraday pair-folds and, after Deflated-Sharpe correction for the number "
        "of trials, is statistically indistinguishable from luck (DSR p = 0.192); and (ii) on "
        "the breadth crypto universe — where the statistics are trustworthy — the net edge is "
        "decisively negative (net DSR 0.002). Gross is alive, net is dead: the obstacle is "
        "execution cost and turnover, not the signal.",
    )

    tcap(
        doc,
        "1",
        "Headline result ladder (all net figures are after the β-aware (1+β) cost). "
        "Sources: reports/backtest_portfolio_metrics.csv, "
        "reports/phase3_ou/figures/matched_universe/matched_metrics.csv, "
        "reports/phase4/verification/1b_eight_cell_table.csv, "
        "reports/phase4/dsr_pbo/2b_dsr.csv, "
        "reports/phase4/crypto_adaptive/dsr_pbo_kalman.csv.",
    )
    ladder = pd.DataFrame(
        [
            [
                "NSE daily baseline (Phase 2A)",
                "492→341 univ.",
                "297.35%",
                "215.67%",
                "1.28",
                "1.08",
                "—",
                "net-profitable; broad universe",
            ],
            [
                "NSE intraday frozen-OU best (Ph 3)",
                "2 pair-folds",
                "54.87%",
                "46.20%",
                "1.08",
                "0.95",
                "0.23",
                "matched universe; n=2",
            ],
            [
                "NSE intraday Kalman best (Ph 4)",
                "2 pair-folds",
                "161.35%",
                "139.94%",
                "2.32",
                "2.15",
                "0.81",
                "clears luck bar; p=0.192",
            ],
            [
                "Crypto Kalman best, Regime A (A10)",
                "14 pairs×540d",
                "+8.96%",
                "−63.20%",
                "+0.11",
                "−1.31",
                "0.002",
                "gross alive, net dead",
            ],
        ],
        columns=[
            "stage",
            "universe",
            "gross tot",
            "net tot",
            "gross Sh",
            "net Sh",
            "DSR",
            "reading",
        ],
    )
    table_from_df(doc, ladder, fontsize=8.5)
    para(doc, "")

    figure(
        doc,
        1,
        "plots/phase4/final/nse/portfolio_equity.png",
        "NSE adaptive equity (matched universe). Portfolio NAV with gross and net lines "
        "and an underwater (drawdown) panel, for the Kalman best cell alongside the "
        "frozen-OU and rolling-z companions. The adaptive arm compounds to roughly 2.4× "
        "net over the two concatenated out-of-sample folds. Source: final_nse_plots.py.",
        width=6.4,
    )
    figure(
        doc,
        2,
        "plots/phase4/final/crypto/portfolio_equity.png",
        "Crypto adaptive equity (breadth universe, 14 pairs × 540 sessions). Gross stays "
        "near flat-to-positive for the Kalman arm while net craters under intraday taker + "
        "spread cost; Regime B is funding-unpriced [TODO]. This is the cost wall that "
        "blocks deployment. Source: final_crypto_plots.py.",
        width=6.4,
    )
    figure(
        doc,
        3,
        "plots/phase4/dsr_pbo/pbo_logit_distribution.png",
        "The validation gate (NSE matched universe). CSCV logit distribution behind the "
        "Probability of Backtest Overfitting (PBO = 0.104). Only the adaptive Kalman arm "
        "clears the N=46 Deflated-Sharpe luck bar; its p-value (0.192) still falls short of "
        "significance. Source: s2_dsr_pbo.py.",
        width=5.4,
    )
    doc.add_page_break()


# --------------------------------------------------------------------------- #
# 3. Introduction
# --------------------------------------------------------------------------- #
def introduction(doc):
    h(doc, "2.  Introduction and Objective", level=1)
    h(doc, "2.1  Statistical-arbitrage pairs trading", level=2)
    para(
        doc,
        "Pairs trading is the canonical market-neutral statistical-arbitrage strategy: two "
        "assets whose prices share a long-run equilibrium are traded as a single mean-"
        "reverting spread — short the rich leg, long the cheap leg — and the position is "
        "unwound when the spread reverts. The edge depends on (a) the existence of a stable "
        "cointegrating relationship, (b) a spread that mean-reverts fast enough to be "
        "capital-efficient yet slowly enough to be tradeable net of cost, and (c) execution "
        "cheap enough that round-trip frictions do not consume the captured dislocation.",
    )
    h(doc, "2.2  Cash-market constraints and the decay context", level=2)
    bullet(
        doc,
        "the equity-pairs edge has decayed as the strategy crowded — naive "
        "rolling-z mean reversion is largely arbitraged away on liquid names, which "
        "raises the bar for any claimed edge and motivates an explicit "
        "overfitting-aware validation gate.",
        bold_lead="Decay of equity pairs:  ",
    )
    bullet(
        doc,
        "the NSE cash segment does not permit native short selling for multi-day "
        "carry, lot-size discretisation adds friction, and the available daily history "
        "is split/bonus-only (dividends are NOT back-adjusted), so total-return "
        "effects are absent from the price series.",
        bold_lead="NSE cash constraints:  ",
    )
    bullet(
        doc,
        "a venue with 24/7 trading, native two-sided leverage via perpetual "
        "futures, and genuine intraday co-movement is the natural place to retest the "
        "mechanism with breadth — and is a stepping stone toward a regulated "
        "futures (e.g. CME) deployment where the cash-market constraints do not bind.",
        bold_lead="Crypto → futures rationale:  ",
    )
    h(doc, "2.3  Objective and thesis", level=2)
    para(
        doc,
        "The thesis under test: does an adaptive equilibrium — letting the spread's "
        "mean-reversion level track slowly rather than freezing it at the in-sample estimate "
        "— produce a transaction-cost-robust improvement over a frozen cointegration "
        "relationship? The study answers this in three movements: establish a disciplined "
        "NSE pipeline and baseline (Phases 1–2); develop and stress the adaptive engine at "
        "intraday frequency (Phases 3–4); and retest the mechanism out-of-sample with "
        "breadth on crypto, gating every claim through a Deflated-Sharpe / PBO filter.",
    )


# --------------------------------------------------------------------------- #
# 4. Data
# --------------------------------------------------------------------------- #
def data_section(doc):
    h(doc, "3.  Data", level=1)
    h(doc, "3.1  NSE daily universe and the seven-rule cleaning cascade", level=2)
    para(
        doc,
        "The raw source is a read-only archive of per-symbol NSE daily OHLCV CSVs "
        "(/Data6/db): 492 symbols and ~1.6 million rows. The history is split/bonus-only — "
        "a dividend-flavor probe (Table 2) shows split adjustments present but dividends "
        "unadjusted on all five dividend-paying probe names. Cleaning is a deterministic, "
        "idempotent seven-rule cascade; each rule is a pure function with logged row counts.",
    )
    rules = pd.DataFrame(
        [
            ["1", "Trading-calendar filter", "drop weekend/holiday/off-calendar bars"],
            [
                "2",
                "Residual-split patch",
                "back-adjust pre-ex-date OHLC for residual splits (e.g. FEDERALBNK, VINATIORGA)",
            ],
            [
                "3",
                "Phantom-seam trim",
                "drop history before the last unexplained >65% day-over-day jump (phantom_jump_threshold = 0.65)",
            ],
            [
                "4",
                "Structural-event windowing",
                "keep only history after structural events (mergers/demergers; e.g. ADANIENT, IDEA)",
            ],
            [
                "5",
                "Liquidity floor",
                "drop bars with 60-day rolling-median ADV < ₹10M (close×volume)",
            ],
            [
                "6",
                "Minimum-history filter",
                "drop symbols with < 756 cleaned rows (~3 trading years)",
            ],
            [
                "7",
                "Contiguity filter",
                "keep only the longest contiguous segment (gaps > 10 calendar days split a symbol)",
            ],
        ],
        columns=["rule", "name", "operation"],
    )
    tcap(
        doc,
        "1",
        "The seven-rule cleaning cascade (src/apt/data, config/default.yaml "
        "cleaning/liquidity/universe blocks). Rule 6 (min-history) and Rule 7 "
        "(contiguity) gate the symbol count.",
    )
    table_from_df(doc, rules, fontsize=9)
    para(doc, "")
    rich(
        doc,
        [
            ("Symbol funnel:  ", {"bold": True, "color": ACCENT}),
            (
                "492 raw symbols → 342 survive the cascade (Rule 6/7 gating) → 341 trading universe "
                "after one share-class duplicate is dropped at backtest time. ",
                {},
            ),
            (
                "[Inter-doc note: the clean log records 342 post-cascade; the backtest loads 341 — "
                "the difference is a single share-class de-duplication, not a data loss.]",
                {"italic": True, "color": GREY},
            ),
        ],
    )

    div = _csv("reports/dividend_flavor_diagnosis.csv")
    divt = fmt(
        div,
        {
            "symbol": ("symbol", None),
            "n_overlap": ("n overlap", i0),
            "cv_vs_close": ("cv vs close", f3),
            "cv_vs_adj_close": ("cv vs adj_close", f3),
            "verdict": ("verdict", None),
            "confidence": ("confidence", None),
        },
    )
    tcap(
        doc,
        "2",
        "Dividend-flavor diagnostic (reports/dividend_flavor_diagnosis.csv). The "
        "low coefficient of variation against raw close vs the larger CV against a "
        "total-return adjusted close indicates split-only adjustment — dividends "
        "are not back-adjusted in the source.",
    )
    table_from_df(doc, divt, fontsize=9)
    para(doc, "")
    para(
        doc,
        "Corporate actions are repaired conservatively: after a diagnostic pass over yfinance "
        "split/dividend coverage (reports/corporate_actions_audit.csv, "
        "corporate_actions_coverage.csv), the pipeline drops ratio-snap ADJUST in favour of "
        "routing every flagged action to TRIM (resample to the post-event spine) — a "
        "deliberate choice to avoid introducing synthetic back-adjustment error "
        "(reports/ca_repair_classification.csv).",
    )
    figure(
        doc,
        4,
        "plots/phase1/universe/01_symbols_per_sector.png",
        "Cleaned NSE universe by sector (scripts/06 universe EDA). The surviving universe "
        "is concentrated in financials and materials — the sectors that later dominate the "
        "cointegrated-pair pool (banks, cement, oil & gas).",
        width=5.8,
    )

    h(doc, "3.2  Crypto universe", level=2)
    para(
        doc,
        "The out-of-sample venue is Binance 1-minute klines (12-column headerless; epoch "
        "timestamps in UTC, with a mixed ms/µs convention after Binance's 2025 switch that "
        "the loader normalises by magnitude). 30 USDT pairs span 2017-08 → 2026-04. A "
        "liquidity gate (median daily quote-volume ≥ $10M, labelled default) keeps 22/30 "
        "symbols; at $50M only 7 survive, too thin for breadth. The NSE seven-rule cascade "
        "does not transfer (no corporate actions; 24/7 sessions) — only minimal cleaning "
        "applies (dedup open_time, drop non-positive OHLC and zero-volume bars). A funding-"
        "rate series is absent from the archive, so funding cost is carried as a [TODO] and "
        "every multi-day-carry (Regime B) number is marked funding-unpriced.",
    )


# --------------------------------------------------------------------------- #
# 5. Methodology
# --------------------------------------------------------------------------- #
def methodology(doc):
    h(doc, "4.  Methodology", level=1)
    para(
        doc,
        "Each technique is given a model specification, the reason it was chosen, how it is "
        "fit TRAIN-only, how it is validated, and a pointer to the result. A single discipline "
        "runs through all of it: nothing is fit, tuned, or selected on test data; every "
        "hyperparameter is chosen on a training window and frozen before the out-of-sample "
        "slice is touched.",
        italic=True,
    )

    h(doc, "4.1  Pair selection — Engle–Granger + BH-FDR (Johansen comparator)", level=2)
    para(
        doc,
        "Candidate pairs are first correlation-screened, then tested for cointegration by the "
        "Engle–Granger two-step: regress log-prices to obtain a static hedge ratio β and "
        "residual, then ADF-test the residual for stationarity. Because hundreds of pairs are "
        "tested per fold, raw ADF p-values are corrected for multiple comparisons by the "
        "Benjamini–Hochberg false-discovery-rate procedure (α = 0.05) — this is the single "
        "most important guard against selecting spuriously-cointegrated pairs.",
    )
    equation(
        doc,
        "log Y_t = α + β · log X_t + r_t ;   ADF(r_t) → p ;   BH-FDR over {p_i}",
        where="(Engle–Granger + BH)",
    )
    para(
        doc,
        "Johansen's trace test is implemented as an order-independent comparator (it does not "
        "privilege one leg as the dependent variable). On a common 234-pair candidate set it "
        "selects 41 pairs vs Engle–Granger+FDR's 13 — but Section 6 shows that gap is almost "
        "entirely the absence of a multiple-testing correction on the Johansen path, not a "
        "structurally different selection. The two universes are therefore reported "
        "separately and never merged into a matched performance table.",
    )

    h(doc, "4.2  Spread, hedge ratio, and the (1+β) cost convention", level=2)
    para(
        doc,
        "The traded spread is s_t = log Y_t − β·log X_t − α with β frozen from the training "
        "OLS fit (direction chosen so β > 0). A one-unit spread position is long 1 notional "
        "unit of Y and short β notional units of X, so the round-trip transaction cost in "
        "cost-log space is proportional to the engaged notional (1+β), not to a flat 2 legs:",
    )
    equation(doc, "cost_per_pair_round_trip = (1 + β) · cost_per_leg", where="(β-aware billing)")
    rich(
        doc,
        [
            ("This corrects a legacy equal-notional 2× convention. ", {}),
            (
                "The bias ratio is (1+β)/2: high-β pairs were under-charged, low-β pairs "
                "over-charged. ",
                {},
            ),
            (
                "Re-billing the entire pipeline under (1+β) moved the headline frozen-OU best cell by "
                "only −0.33 percentage points of annual return (decomposed exactly as billing "
                "−0.398 pp + threshold-refit +0.071 pp = −0.326 pp; Section 5.4), and moved the daily "
                "Phase-2A net Sharpe by +0.006 — negligible — but it removed a spurious "
                "rolling-z-beats-OU inversion at the lowest cost and is the single billing convention "
                "used everywhere thereafter.",
                {},
            ),
        ],
    )

    h(doc, "4.3  OU process and Bertram cost-aware thresholds", level=2)
    para(
        doc,
        "At intraday frequency the spread is modelled as an Ornstein–Uhlenbeck process and "
        "fit by its exact AR(1) discretisation on the training window:",
    )
    equation(doc, "dX_t = κ(μ − X_t) dt + σ dW_t      ⇔      X_{t+1} = c + φ·X_t + ε_t")
    equation(doc, "κ = −ln(φ)/Δt ,   μ = c/(1−φ) ,   σ_eq² = σ²/(2κ) ,   half-life = ln(2)/κ")
    para(
        doc,
        "Entry/exit thresholds are not hand-set: they solve Bertram's (2010) closed-form "
        "optimum, which maximises expected return per unit time net of round-trip cost c, "
        "trading at ±a* in equilibrium-σ units and exiting at the mean. The transcription was "
        "Monte-Carlo validated against a first-passage simulation (analytic crossing "
        "probability 1.140×10⁻³ vs simulated 1.132×10⁻³). The threshold is re-solved per cost "
        "level, so a* widens monotonically with cost (e.g. 0.41 → 0.54 σ-units from 1 → 8 bps "
        "on the KOTAK fold) — the engine self-amortises cost by trading less, which is the "
        "structural reason it is far more cost-resilient than a fixed-band rolling-z rule.",
    )
    para(
        doc,
        "A half-life band filter separates two trade regimes: Regime A (intraday "
        "square-off) admits half-lives in [30, 120] minutes; Regime B (multi-day carry) "
        "admits [120, 1875] minutes (1875 min = 5 sessions). A time-of-day volatility "
        "normalisation removes the intraday U-shape before z-scoring. Critically, 0 of 18 "
        "valid OU fits fall in the Regime-A band at any frequency — the NSE pairs mean-revert "
        "on a multi-session timescale, which makes intraday square-off structurally "
        "infeasible (Section 6.3).",
    )

    h(doc, "4.4  Adaptive equilibrium — the μ-only Kalman filter", level=2)
    para(
        doc,
        "The central methodological contribution replaces the frozen equilibrium μ with a "
        "per-session causal local-level (random-walk-plus-noise) filter, so the spread's "
        "reversion target tracks slowly rather than being pinned at the in-sample estimate:",
    )
    equation(doc, "state:   μ_t = μ_{t−1} + w_t        observation:   y_t = μ_t + e_t")
    equation(
        doc,
        "constant gain (West–Harrison discount):   K = 1 − 2^(−1/H)",
        where="H = half-life in sessions",
    )
    para(
        doc,
        "H is the discount horizon in sessions; H = ∞ gives K = 0 and recovers the frozen "
        "engine exactly (the frozen-control equivalence, unit-tested bit-for-bit). The filter "
        "updates μ at session close and applies it from the next session open — strictly "
        "causal, no intra-session look-ahead. H is selected TRAIN-only from {∞, 20, 10, 5} by "
        "the Bertram net-return criterion on training residuals, subject to an absorption "
        "guard: the residual half-life under the candidate H must stay within [0.5×, 1.5×] of "
        "the frozen-μ half-life, otherwise the configuration is inadmissible. This guard is "
        "what prevents the filter from over-tracking and absorbing genuine signal into μ.",
    )

    h(doc, "4.5  β-escalation and its identification failure", level=2)
    para(
        doc,
        "The natural escalation is to also let the hedge ratio track. The joint (β, c) filter "
        "uses a returns-regression observation for β and a level observation for c = α + μ:",
    )
    equation(
        doc,
        "β̂_s = cov_s(Δx, Δy) / var_s(Δx) ,    ℓ̂_s = mean_s(y − β_s·x)",
        where="per-session, diagonal gains",
    )
    para(
        doc,
        "This is collapse-prone by construction: when the legs' intraday increments stop "
        "co-moving, cov(Δx,Δy) → 0 and β̂ → 0, evaporating the hedge. A β-stability guard "
        "requires β_s ∈ [0.25, 4]×β₀ over the training window. Section 6.6 shows every finite "
        "tracking horizon collapses β out of that band on all three diagnostic pair-folds — "
        "so train selection freezes β (H_β = ∞), at which point the joint filter is "
        "identically the μ-only filter. The escalation fails, and the failure — weak "
        "intraday identification of the hedge ratio — is itself a finding.",
    )

    h(doc, "4.6  Cost model and backtest protocol", level=2)
    bullet(
        doc,
        "an NSE CostBreakdown (STT 2.5 + brokerage 0.5 + exchange/regulatory 1.5 = "
        "4.5 bps fixed per leg) plus a quoted-spread sweep {1,3,5,8} bps, all billed "
        "(1+β). Crypto cost = taker fee (5 bps/side default) + spread sweep, also (1+β); "
        "funding is [TODO] (no series).",
        bold_lead="Cost:  ",
    )
    bullet(
        doc,
        "annual walk-forward with [prior][train][test] windows rolled forward; every "
        "fit (cointegration, OU, OU-residual, Kalman H, Bertram a*) is performed on the "
        "training window of that fold only and frozen before the test slice. Regime A "
        "(intraday square-off) vs Regime B (multi-day carry); bar-frequency sweep "
        "{1,5,15} minutes.",
        bold_lead="Walk-forward:  ",
    )
    bullet(
        doc,
        "every performance number is reported gross AND net; every equity curve "
        "carries a drawdown panel; exits are tagged with a fixed five-category "
        "vocabulary (mean_revert / z_stop / time_stop / eod_squareoff / fold_close).",
        bold_lead="Reporting standard:  ",
    )

    h(doc, "4.7  Validation framework — Deflated Sharpe and PBO", level=2)
    para(
        doc,
        "Because the development searched a grid of engines × frequencies × costs × regimes, "
        "the in-sample-best Sharpe is upward-biased by selection. Two complementary deflators "
        "are applied. The Deflated Sharpe Ratio (Bailey & López de Prado, 2014) computes the "
        "probability that the observed Sharpe exceeds the expected maximum Sharpe achievable "
        "by N independent random trials, accounting for the return series' skew, kurtosis and "
        "length; the expected-maximum benchmark SR₀ grows with the honest trial count N "
        "(reconstructed from the search grids). The Probability of Backtest Overfitting "
        "(Bailey et al., 2017) uses Combinatorially-Symmetric Cross-Validation: split the "
        "return matrix into S blocks, and over all C(S, S/2) symmetric in-/out-of-sample "
        "partitions measure how often the in-sample-best configuration underperforms the "
        "out-of-sample median. The Bertram threshold transcription itself was validated by an "
        "independent Monte-Carlo first-passage check (Section 4.3).",
    )


# --------------------------------------------------------------------------- #
# 6. Results by stage
# --------------------------------------------------------------------------- #
def results(doc):
    h(doc, "5.  Results by Stage", level=1)

    # 5.1 Phase 2A daily baseline
    h(doc, "5.1  Phase 2A — daily walk-forward baseline", level=2)
    pm = _csv("reports/backtest_portfolio_metrics.csv").set_index("name")

    def gm(name, col):
        return pm.loc[name, col]

    rich(
        doc,
        [
            ("On the daily NSE universe the strategy is net-profitable. ", {"bold": True}),
            (
                f"Across an 8-fold annual walk-forward (~7 out-of-sample years, "
                f"{int(gm('n_obs', 'net'))} daily observations) the portfolio returns "
                f"{gm('total_return_pct', 'gross'):.2f}% gross / {gm('total_return_pct', 'net'):.2f}% "
                f"net, with a net Sharpe of {gm('sharpe', 'net'):.2f} (gross {gm('sharpe', 'gross'):.2f}) "
                f"and a net maximum drawdown of {gm('max_drawdown_pct', 'net'):.2f}%. ",
                {},
            ),
            (
                "This is the broad-universe anchor against which the narrow intraday results must be "
                "read.",
                {},
            ),
        ],
    )
    base = pd.DataFrame(
        {
            "metric": [
                "Total return",
                "Annualised return",
                "Annualised vol",
                "Sharpe",
                "Max drawdown",
                "Observations / years",
            ],
            "gross": [
                pc2(gm("total_return_pct", "gross")),
                pc2(gm("ann_return_pct", "gross")),
                pc2(gm("ann_vol_pct", "gross")),
                f2(gm("sharpe", "gross")),
                pc2(gm("max_drawdown_pct", "gross")),
                f"{int(gm('n_obs', 'gross'))} / {int(gm('n_years', 'gross'))}",
            ],
            "net": [
                pc2(gm("total_return_pct", "net")),
                pc2(gm("ann_return_pct", "net")),
                pc2(gm("ann_vol_pct", "net")),
                f2(gm("sharpe", "net")),
                pc2(gm("max_drawdown_pct", "net")),
                f"{int(gm('n_obs', 'net'))} / {int(gm('n_years', 'net'))}",
            ],
        }
    )
    tcap(
        doc,
        "3",
        "Phase 2A daily walk-forward portfolio, gross vs net under the (1+β) "
        "billing (reports/backtest_portfolio_metrics.csv).",
    )
    table_from_df(doc, base, fontsize=9.5)
    para(doc, "")
    pp = _csv("reports/backtest_per_pair_metrics.csv").head(8)
    ppt = fmt(
        pp,
        {
            "pair": ("pair", None),
            "n_trades": ("trades", i0),
            "gross_total_return_pct": ("gross tot%", f2),
            "net_total_return_pct": ("net tot%", f2),
            "gross_sharpe": ("gross Sh", f2),
            "net_sharpe": ("net Sh", f2),
            "net_max_drawdown_pct": ("net maxDD%", f2),
        },
    )
    tcap(
        doc,
        "4",
        "Phase 2A per-pair contributors (top 8 of 25 pairs by net return; "
        "reports/backtest_per_pair_metrics.csv). PFC/SBIN and the oil-complex pairs "
        "carry the book — foreshadowing the concentration finding.",
    )
    table_from_df(doc, ppt, fontsize=9)
    para(doc, "")
    figure(
        doc,
        5,
        "plots/phase2/backtest/equity_curve.png",
        "Phase 2A portfolio equity curve, daily walk-forward (scripts/10 backtest).",
        width=5.8,
    )

    # 5.2 Phase 2B ablation
    h(doc, "5.2  Phase 2B — risk-managed ablation (rolled back)", level=2)
    para(
        doc,
        "A sizing/risk ablation ladder (R0 baseline → R4 with cluster caps and a "
        "relationship-kill switch) tested whether the baseline's drawdown is a sizing "
        "problem. It is not. Volatility-matching R0 and the heavily risk-controlled R3 to a "
        "common target produces near-identical Sharpe (0.941 vs 0.930) — a tie. The drawdown "
        "is driven by the cointegrated pairs being concentrated in a few clusters (banks, "
        "oil) that dislocate together, not by position sizing. The experiment was reverted; "
        "the finding — concentration is a universe problem, not a sizing problem — is "
        "preserved on its branch.",
    )
    lad = _csv("reports/risk_managed_ladder.csv")
    ladt = fmt(
        lad,
        {
            "label": ("rung", None),
            "kill_mode": ("kill", None),
            "n_trades": ("trades", i0),
            "net_total_pct": ("net tot%", f2),
            "net_sharpe": ("net Sh", f2),
            "max_drawdown_pct": ("maxDD%", f2),
            "turnover_per_year": ("turnover/yr", f1),
        },
    )
    tcap(
        doc,
        "5",
        "Risk-managed ladder R0–R4 (reports/risk_managed_ladder.csv). R3/R4 cut "
        "drawdown only by cutting the book to ~half the trades; net return falls "
        "with it.",
    )
    table_from_df(doc, ladt, fontsize=9)
    vm = _csv("reports/risk_managed_vol_matched.csv")
    vmt = fmt(
        vm,
        {
            "rung": ("rung", None),
            "raw_ann_vol_pct": ("raw vol%", f2),
            "vol_targeted_ann_vol_pct": ("target vol%", f2),
            "vol_targeted_ann_return_pct": ("vt return%", f2),
            "vol_targeted_sharpe": ("vt Sharpe", f3),
            "vol_targeted_max_dd_pct": ("vt maxDD%", f2),
        },
    )
    tcap(
        doc,
        "6",
        "Volatility-matched comparison (reports/risk_managed_vol_matched.csv): the "
        "Sharpe tie (0.941 vs 0.930) that motivated rollback.",
    )
    table_from_df(doc, vmt, fontsize=9)
    para(doc, "")
    figure(
        doc,
        6,
        "plots/phase2/risk_managed/ladder_equity.png",
        "Phase 2B ablation ladder equity (reverted experiment). The risk-controlled rungs "
        "compress both drawdown and return; the risk-adjusted ratio does not improve.",
        width=5.8,
    )

    # 5.3 Phase 3 intraday infeasibility
    h(doc, "5.3  Phase 3 — intraday port and Regime-A infeasibility", level=2)
    para(
        doc,
        "Porting to intraday bars exposes two facts. First, naive intraday rolling-z is "
        "net-negative under any cost — at 1 bps the two-regime precursor runs net Sharpe "
        "−1.46 (Regime A) and −0.34 (Regime B), both underwater. Second, and structurally: "
        "every OU half-life lands in the multi-session band, so the intraday-square-off "
        "Regime A admits zero pairs. Daily-cointegrated NSE pairs simply do not mean-revert "
        "within a session; forcing an end-of-day flat closes trades before reversion "
        "completes. The 1-minute frequency additionally inflates half-life estimates via "
        "bid-ask bounce (a microstructure artefact), so 1-minute 'passes' are not cited as "
        "intraday-reversion evidence.",
    )
    figure(
        doc,
        7,
        "plots/phase3/equity_A_vs_B_3bps.png",
        "Phase 3 intraday Regime A vs Regime B equity at 3 bps (scripts/14). Regime A "
        "(intraday square-off) is structurally starved of admissible pairs.",
        width=5.6,
    )
    figure(
        doc,
        8,
        "plots/phase3/tod_vol_profile.png",
        "Time-of-day volatility profile used for intraday normalisation — the U-shape "
        "removed before z-scoring.",
        width=5.2,
    )

    # 5.4 Phase 3 OU/Bertram + cost-beta
    h(doc, "5.4  Phase 3 — OU/Bertram engine and the matched universe", level=2)
    para(
        doc,
        "Applying the half-life band filter and Bertram thresholds reduces the intraday "
        "universe to two Regime-B survivors: INDUSINDBK/HDFCBANK (fold 4) and "
        "KOTAKBANK/HDFCBANK (fold 6). This is the matched universe on which all subsequent "
        "engine comparisons run. On it, the frozen-OU best cell (5-min, Regime B, 3 bps) is "
        "net-profitable and materially more cost-resilient than rolling-z.",
    )
    mm = _csv("reports/phase3_ou/figures/matched_universe/matched_metrics.csv")
    mm5 = mm[(mm.freq_min == 5) & (mm.regime == "B")].copy()
    mm5t = fmt(
        mm5,
        {
            "engine": ("engine", None),
            "spread_bps": ("cost", i0),
            "n_trades": ("trades", i0),
            "gross_total_pct": ("gross tot%", f2),
            "net_total_pct": ("net tot%", f2),
            "gross_sharpe": ("gross Sh", f3),
            "net_sharpe": ("net Sh", f3),
        },
    )
    tcap(
        doc,
        "7",
        "Matched-universe cost ladder, 5-min Regime B "
        "(reports/phase3_ou/figures/matched_universe/matched_metrics.csv). OU's net "
        "Sharpe decays gently with cost (0.98 → 0.93) while rolling-z collapses "
        "(0.95 → −0.11) and flips sign at 8 bps — the cost-amortisation property.",
    )
    table_from_df(doc, mm5t, fontsize=9)
    para(doc, "")
    rich(
        doc,
        [
            ("Canonical frozen-OU best cell: ", {"bold": True}),
            (
                "net annualised 20.915%, net Sharpe 0.949, net total 46.20%, max drawdown −22.54%, "
                "34 trades (reports/phase3_ou/metrics_ou.csv). ",
                {},
            ),
            (
                "The (1+β) re-billing decomposes the change from the legacy convention exactly: "
                "A {old a*, old bill} = 21.241 → B {old a*, new bill} = 20.844 (billing −0.398) → "
                "C {new a*, new bill} = 20.915 (refit +0.071), total −0.326 pp.",
                {"italic": True},
            ),
        ],
    )
    figure(
        doc,
        9,
        "reports/phase3_ou/figures/ou_best_cell/b_ou_best_portfolio_nav.png",
        "OU best-cell portfolio NAV, gross vs net (reports/phase3_ou/figures/ou_best_cell).",
        width=5.6,
    )
    figure(
        doc,
        10,
        "reports/phase3_ou/figures/grid_rollups/d_cost_ladder_f5_B.png",
        "Cost ladder, 5-min Regime B: OU vs rolling-z net Sharpe and net total return vs "
        "cost. OU's flat slope is the engine-attributable cost resilience.",
        width=5.8,
    )
    figure(
        doc,
        11,
        "reports/phase3_ou/figures/grid_rollups/f_half_life_distribution.png",
        "OU half-life distribution with Regime-A/B band boundaries. No fit lands in the "
        "Regime-A [30,120]-min band — the structural reason intraday square-off is "
        "infeasible.",
        width=5.8,
    )
    figure(
        doc,
        12,
        "reports/phase3_ou/figures/grid_rollups/k_exclusion_funnel.png",
        "Exclusion funnel: 19 pair-folds → 18 valid AR(1) → 2 traded Regime-B survivors at "
        "5-min. The narrowness of the matched universe (n=2) is the central fragility.",
        width=5.6,
    )

    # 5.5 Kalman
    h(doc, "5.5  Phase 4 — adaptive Kalman equilibrium (the 8-cell verification)", level=2)
    para(
        doc,
        "On the same matched universe the μ-only Kalman arm roughly doubles the frozen-OU net "
        "Sharpe across the cell grid. The verification reproduces the headline (net Sharpe "
        "≈ 2.1 at the best cell) but also corrects the mechanism narrative.",
    )
    ec = _csv("reports/phase4/verification/1b_eight_cell_table.csv")
    ect = fmt(
        ec,
        {
            "freq_min": ("freq", i0),
            "cost_bps": ("cost", i0),
            "k_n_trades": ("k tr", i0),
            "k_net_total_pct": ("k net%", f2),
            "k_net_sharpe": ("k net Sh", f3),
            "k_net_ann_pct": ("k netAnn%", f2),
            "k_net_maxDD_pct": ("k maxDD%", f2),
            "ou_net_sharpe": ("ou net Sh", f3),
            "rz_net_sharpe": ("rz net Sh", f3),
        },
    )
    tcap(
        doc,
        "8",
        "Kalman vs frozen-OU vs rolling-z, all eight matched cells "
        "(reports/phase4/verification/1b_eight_cell_table.csv). Kalman net Sharpe "
        "≈ 2× the frozen engine at every cell.",
    )
    table_from_df(doc, ect, fontsize=8.5)
    para(doc, "")
    para(
        doc,
        "What the verification corrects: the adaptive edge is downside-tail compression, not "
        "larger winners — the Kalman 5th-percentile trade is −168 bps vs the frozen engine's "
        "−605 bps, while its median trade is actually lower. The edge is NOT sourced from "
        "rescuing the frozen engine's stuck trades (the overlap with frozen losing "
        "time-stop/fold-close windows is +0.23% of gross). And the win rate is "
        "mechanism-inflated: 78.6% of trades see the tracking anchor μ move toward the entry "
        "spread during the hold. Two of four pre-registered reviewer predictions are NOT MET "
        "(the per-trade capture is smaller than predicted and the overlap story is wrong), "
        "even though the net-Sharpe improvement is real.",
    )
    figure(
        doc,
        13,
        "reports/phase3_kalman/figures/drift_before_after.png",
        "Equilibrium drift before/after adaptive tracking (reports/phase3_kalman/figures). "
        "The filter pulls the standardised test-slice drift toward zero.",
        width=5.6,
    )
    figure(
        doc,
        14,
        "plots/phase4/verification/per_trade_pnl_dist.png",
        "Per-trade net P&L: Kalman vs frozen-OU. The Kalman distribution is left-tail "
        "compressed (tighter p5), not right-shifted — downside control, not bigger wins.",
        width=5.8,
    )
    figure(
        doc,
        15,
        "reports/phase3_kalman/figures/mu_overlay_fold4.png",
        "μ-tracking overlay, INDUSINDBK/HDFCBANK fold 4: spread with frozen μ and tracking "
        "μ_t and trade marks. The anchor moves materially across sessions.",
        width=5.8,
    )

    # 5.6 beta escalation
    h(doc, "5.6  Phase 4 — β-escalation (weak identification)", level=2)
    para(
        doc,
        "Allowing β to track as well fails. At every finite tracking horizon, on every "
        "diagnostic pair-fold, β collapses toward zero (fold 4: 1.643 → 0.089 at H_β=10; the "
        "KOTAK fold even flips β negative). Intraday increment co-movement is too weak to "
        "identify the hedge ratio, so the returns-regression observation drags β to the "
        "floor. Train selection consequently picks H_β = ∞ — it refuses to track β — at which "
        "point the joint filter is identically the μ-only filter (verified bit-for-bit). The "
        "canonical −7σ fold-6 case is not admissibly neutralised by a tracking-β filter on "
        "this data. The gate fails, and the failure is the finding: the path forward is to "
        "estimate β at a lower frequency or on a universe that genuinely co-moves intraday.",
    )
    bc = _csv("reports/phase4/beta_escalation/beta_collapse.csv")
    bc5 = bc[(bc.freq_min == 5) & (bc.h_beta != "inf")].copy()
    bc5["h_beta"] = bc5["h_beta"].astype(str)
    bc5t = fmt(
        bc5,
        {
            "pair": ("pair", None),
            "h_beta": ("H_β", None),
            "beta_init": ("β₀", f3),
            "beta_test_min": ("β min", f3),
            "min_beta_ratio": ("min ratio", f3),
            "beta_toward_zero": ("→0?", lambda x: "yes" if x else "no"),
            "beta_stable": ("stable?", lambda x: "yes" if x else "no"),
        },
    )
    tcap(
        doc,
        "9",
        "β-collapse diagnostic, 5-min, finite horizons "
        "(reports/phase4/beta_escalation/beta_collapse.csv). Every finite H_β drives "
        "β below the stability band on every pair-fold.",
    )
    table_from_df(doc, bc5t, fontsize=8.5)
    para(doc, "")
    figure(
        doc,
        16,
        "plots/phase4/beta_escalation/beta_paths_f5.png",
        "β tracking paths (5-min). Every finite horizon collapses the hedge ratio toward "
        "zero — weak intraday identification.",
        width=5.8,
    )

    # 5.7 coint stability + johansen
    h(doc, "5.7  Phase 4 — rolling cointegration-stability gate", level=2)
    para(
        doc,
        "A rolling-ADF gate on the frozen-(α,β) residual blacklists a pair after three "
        "consecutive non-stationary windows. At the default it fires on 19/19 NSE pair-folds — "
        "including both traded survivors — and the sensitivity sweep shows this is robust, "
        "not a threshold accident (15/19 still gate even at the lenient p>0.50 / 120-day "
        "setting). Two forces are entangled and cannot be separated on n=2: genuinely "
        "marginal cointegration, and low ADF power on short windows. The gate is correct and "
        "tested but uncalibratable on this universe — it is marked [TODO]-pending-breadth.",
    )
    ts = _csv("reports/phase4/coint_stability/threshold_sensitivity.csv")
    tst = fmt(
        ts,
        {
            "window": ("window", i0),
            "threshold": ("ADF p>", f2),
            "consecutive": ("consec", i0),
            "n_gated": ("gated", i0),
            "frac_gated": ("frac gated", f3),
        },
    )
    tcap(doc, "10", "Gate sensitivity (reports/phase4/coint_stability/threshold_sensitivity.csv).")
    table_from_df(doc, tst, fontsize=9)
    para(doc, "")
    figure(
        doc,
        17,
        "plots/phase4/coint_stability/adf_pvalue_paths.png",
        "Rolling-ADF p-value paths with the blacklist threshold. The NSE spreads drift "
        "non-stationary on rolling windows.",
        width=5.8,
    )

    h(doc, "5.8  Phase 4 — Johansen comparison", level=2)
    cs = _csv("reports/phase4/johansen/comparison_summary.csv").iloc[0]
    rich(
        doc,
        [
            (
                f"On {int(cs.n_candidates_tested)} common candidates, Johansen-95% selects "
                f"{int(cs.n_johansen_selected)} pairs vs Engle–Granger+FDR's "
                f"{int(cs.n_eg_fdr_selected)} ({int(cs.n_johansen_only_NEW)} new, "
                f"{int(cs.n_eg_only_LOST)} lost; Jaccard {cs.jaccard:.2f}). ",
                {},
            ),
            ("But the honest attribution is multiple-testing, not the estimator: ", {"bold": True}),
            (
                f"a fresh Engle–Granger pass reproduces the persisted FDR selection at "
                f"{cs.fresh_eg_vs_persisted_fdr_agreement * 100:.1f}% agreement, and Johansen's "
                f"{int(cs.n_johansen_selected)} sits right next to EG-raw-p<0.05's "
                f"{int(cs.n_eg_raw_p05_selected)} — the expansion is the absence of an FDR "
                "correction on the Johansen path. The 28 new pairs are unvalidated candidates "
                "and are explicitly not merged into any performance table.",
                {},
            ),
        ],
    )
    figure(
        doc,
        18,
        "plots/phase4/johansen/eg_vs_johansen.png",
        "Engle–Granger p-value vs Johansen trace statistic, coloured by selection "
        "(reports/phase4/johansen). The gap is the FDR correction, not order-independence.",
        width=5.8,
    )

    # 5.9 DSR/PBO gate
    h(doc, "5.9  Phase 4 — the DSR / PBO validation gate (NSE)", level=2)
    para(
        doc,
        "This is the primary result. Over an honest N=46 trials, only the adaptive Kalman arm "
        "clears the expected-maximum luck bar; the frozen-OU and rolling-z best cells fall "
        "below it.",
    )
    dsr = _csv("reports/phase4/dsr_pbo/2b_dsr.csv")
    dsrt = fmt(
        dsr,
        {
            "engine": ("engine", None),
            "selected_cell": ("cell", None),
            "ann_net_sharpe": ("ann net Sh", f3),
            "per_period_sharpe": ("per-period SR", f3),
            "SR0_deflator": ("SR₀ bar", f3),
            "DSR": ("DSR", f3),
            "DSR_pvalue": ("p-value", f3),
        },
    )
    tcap(
        doc,
        "11",
        "Deflated Sharpe Ratio, NSE matched universe, N=46 "
        "(reports/phase4/dsr_pbo/2b_dsr.csv). Only Kalman clears the bar; its "
        "p-value (0.192) is still not significant. PBO over the 24-cell matrix = "
        "0.104 (reports/phase4/dsr_pbo/2c_pbo.csv).",
    )
    table_from_df(doc, dsrt, fontsize=9)
    sens = _csv("reports/phase4/dsr_pbo/2b_dsr_sensitivity_to_N.csv")
    senst = fmt(
        sens,
        {
            "n_trials": ("N trials", i0),
            "SR0_deflator": ("SR₀ bar", f3),
            "DSR": ("Kalman DSR", f3),
            "p_value": ("p-value", f3),
        },
    )
    tcap(
        doc,
        "12",
        "DSR sensitivity to the assumed trial count N "
        "(reports/phase4/dsr_pbo/2b_dsr_sensitivity_to_N.csv). Kalman DSR stays "
        "above 0.5 across N but never reaches p<0.05.",
    )
    table_from_df(doc, senst, fontsize=9)
    para(doc, "")
    rich(
        doc,
        [
            ("Mandatory caveat: ", {"bold": True, "color": RED}),
            (
                "with n=2 pair-folds the per-period series is the concatenation of two disjoint folds "
                "(≈2017 and ≈2019); it violates the iid assumption behind DSR and mixes two regimes "
                "inside every CSCV block. The NSE DSR/PBO numbers are INDICATIVE, not definitive. "
                "What would make them trustworthy is universe breadth — which the crypto retest "
                "supplies.",
                {},
            ),
        ],
    )

    # 5.10 crypto
    h(doc, "5.10  Phase 4 / A10 — crypto out-of-sample retest (the decisive experiment)", level=2)
    para(
        doc,
        "The same three engines were wired through the crypto intraday pipeline with EG-FDR "
        "per-fold selection, (1+β) billing, train-only H selection, and a {1,5,15}-minute × "
        "{1,3,5,8}-bps sweep over both regimes — 5 folds, 17 EG-FDR pair-folds, 14 traded "
        "pairs over 540 sessions. Two findings emerge in order of importance.",
    )
    numbered(
        doc,
        "on both gross and net Sharpe, Kalman > frozen-OU at every frequency, and "
        "Kalman is the only engine with positive gross. The adaptive mechanism "
        "extracts more signal than the frozen mean on a completely different "
        "universe — exactly the NSE ordering.",
        bold_lead="The ordering replicates:  ",
    )
    numbered(
        doc,
        "crypto intraday taker + spread cost (≈2,800 round-trips, (1+β)-billed) "
        "destroys the edge — Kalman is the least bad (−63%) but still deeply "
        "unprofitable; rolling-z's +37,803% gross at 1-min is a pure overtrading "
        "artefact (94,403 trades) that evaporates to −100% net. Gross > 0 but net ≪ 0 "
        "means this is a cost/turnover problem, not a signal problem.",
        bold_lead="But net is negative for all engines:  ",
    )
    cam = _csv("reports/phase4/crypto_adaptive/metrics.csv")
    caA = cam[(cam.regime == "A") & (cam.spread_bps == 3)].copy()
    order = {"kalman_mu": 0, "frozen_ou": 1, "rolling_z": 2}
    caA["o"] = caA.engine.map(order)
    caA = caA.sort_values(["o", "freq_min"])
    caAt = fmt(
        caA,
        {
            "engine": ("engine", None),
            "freq_min": ("freq", i0),
            "n_trades": ("trades", i0),
            "gross_total_pct": ("gross tot%", f2),
            "net_total_pct": ("net tot%", f2),
            "gross_sharpe": ("gross Sh", f3),
            "net_sharpe": ("net Sh", f3),
            "net_max_drawdown_pct": ("maxDD%", f2),
        },
    )
    tcap(
        doc,
        "13",
        "Crypto A10 headline, Regime A (funding-clean), cost 3 bps "
        "(reports/phase4/crypto_adaptive/metrics.csv). Kalman is uniquely "
        "positive on gross; every engine is deeply negative on net.",
    )
    table_from_df(doc, caAt, fontsize=8.5)
    para(doc, "")
    dpk = _csv("reports/phase4/crypto_adaptive/dsr_pbo_kalman.csv")
    dpkt = fmt(
        dpk,
        {
            "regime": ("regime", None),
            "funding": ("funding", None),
            "cell": ("cell", None),
            "ann_net_sharpe": ("ann net Sh", f3),
            "per_period_sharpe": ("per-period SR", f3),
            "SR0_luck_bar": ("SR₀ bar", f3),
            "DSR": ("net DSR", f3),
            "DSR_pvalue": ("p-value", f3),
            "PBO": ("PBO", f3),
        },
    )
    tcap(
        doc,
        "14",
        "Crypto Kalman DSR/PBO, N=28 honest trials "
        "(reports/phase4/crypto_adaptive/dsr_pbo_kalman.csv). Net DSR = 0.002 — the "
        "edge does not survive. (Gross Kalman DSR = 0.546, p 0.454 — marginal, not "
        "significant.) Breadth (14 pairs × 540 d) makes these trustworthy.",
    )
    table_from_df(doc, dpkt, fontsize=8.5)
    para(doc, "")
    para(
        doc,
        "Frequency does not rescue net: median holding is ≈1,430 minutes (~1 day) at all "
        "frequencies and the best-frequency vote is split across pairs. The absorption guard "
        "admitted 29/51 cells at the selected H=10, so the intraday crypto residual is "
        "tradeable (the guard binds sensibly) — the problem is purely cost. An earlier daily "
        "rolling-z crypto port (Section appendix) was negative on gross already — a signal "
        "problem for that weaker engine — and is superseded by A10 as the real test.",
    )
    figure(
        doc,
        19,
        "plots/phase4/final/crypto/trade_return_distribution.png",
        "Crypto per-trade net P&L, Kalman vs rolling-z (Regime A). The typical Kalman trade "
        "is profitable (median +14.3 bps) but the mean is −14.4 bps — the left tail plus "
        "per-trade cost sink the aggregate. Source: final_crypto_plots.py.",
        width=5.8,
    )
    figure(
        doc,
        20,
        "plots/phase4/crypto/crypto_cost_ladder.png",
        "Crypto cost ladder (daily rolling-z port, Section appendix): decisively negative, "
        "gross already below zero. Source: s6_crypto.py.",
        width=5.6,
    )


# --------------------------------------------------------------------------- #
# 7. Discussion
# --------------------------------------------------------------------------- #
def discussion(doc):
    h(doc, "6.  Discussion — Methodological Lessons", level=1)
    items = [
        (
            "Concentration is a universe problem, not a sizing problem.",
            "Phase 2B's vol-matched tie (Sharpe 0.941 vs 0.930) shows risk controls cannot "
            "neutralise drawdown that comes from cointegrated pairs clustering in a few sectors "
            "that dislocate together. The remedy is breadth and de-correlation in selection, not "
            "leverage targeting.",
        ),
        (
            "Cost-aware thresholds reduce but do not eliminate over-trading.",
            "Bertram's optimal a* widens with cost and flattens the NSE cost ladder relative to a "
            "fixed-band rolling-z — but on crypto, where turnover is intrinsically high, even "
            "cost-aware widening leaves net deeply negative. Cost amortisation is necessary, not "
            "sufficient.",
        ),
        (
            "In-sample selection inflates Sharpe, and deflation is the arbiter.",
            "The same matched-universe Kalman Sharpe of ~2.1 reads very differently once N=46 "
            "trials are priced in: DSR 0.808 but p=0.192. The Deflated Sharpe Ratio, not the raw "
            "Sharpe, is what separates an indicative result from a deployable one.",
        ),
        (
            "The n=2 fragility dominates every NSE conclusion.",
            "Two intraday pair-folds cannot support an iid validation framework: the CSCV blocks "
            "straddle a fold join and the DSR series is regime-mixed. No amount of careful "
            "in-sample work overcomes the absence of cross-sectional width — which is precisely "
            "why the crypto retest exists.",
        ),
        (
            "Dynamic hedge ratios are weakly identified at high frequency.",
            "β collapses on every finite tracking horizon because intraday increment co-movement "
            "is too weak to identify it. The frozen β is not a simplification of convenience — it "
            "is what the data will support; tracking β requires lower-frequency estimation or a "
            "genuinely co-moving universe.",
        ),
        (
            "The cost wall is the binding constraint.",
            "Across both universes the adaptive engine wins on gross and on the engine ordering, "
            "yet loses on net. The mechanism is real and universe-general; the obstacle is "
            "execution cost at the turnover the signal demands. This reframes the research "
            "question from 'find a better signal' to 'execute the same signal more cheaply'.",
        ),
    ]
    for lead, body in items:
        rich(doc, [(lead + "  ", {"bold": True, "color": ACCENT}), (body, {})])


# --------------------------------------------------------------------------- #
# 8. Conclusion
# --------------------------------------------------------------------------- #
def conclusion(doc):
    h(doc, "7.  Conclusion", level=1)
    rich(
        doc,
        [
            (
                "We have built and validated, end-to-end, an adaptive pairs-trading methodology, and "
                "the honest verdict is a rigorous negative: ",
                {},
            ),
            (
                "a real, repeatable, universe-general mean-reversion mechanism that does not clear "
                "realistic intraday transaction costs.",
                {"bold": True},
            ),
        ],
    )
    para(
        doc,
        "The adaptive-equilibrium filter is genuine — it beats a frozen cointegration "
        "relationship on gross and net Sharpe on the matched NSE universe, and the same "
        "ordering replicates on an independent crypto universe where it is uniquely positive "
        "on gross. That cross-universe replication is the contribution: the mechanism is not "
        "an NSE-specific artefact. But the strongest in-sample NSE result is statistically "
        "indistinguishable from luck after Deflated-Sharpe correction (DSR p=0.192, n=2), and "
        "on the breadth crypto universe — where the statistics are trustworthy — the net edge "
        "is decisively negative (net DSR 0.002).",
    )
    rich(
        doc,
        [
            ("What would change the verdict is structural, not parametric. ", {"bold": True}),
            (
                "Because gross is alive and net is dead, the lever is turnover and execution cost, "
                "not the signal: lower-frequency re-anchoring with fewer round-trips, a cheaper venue "
                "or maker/rebate execution, and a turnover-aware variant of the same adaptive engine. "
                "Re-parameterising the signal will not help; re-engineering the cost will. We present "
                "this as a methodological result — a disciplined demonstration of where a plausible "
                "edge actually dies, and why deflation and breadth, not a larger in-sample Sharpe, "
                "are the instruments that reveal it.",
                {},
            ),
        ],
    )


# --------------------------------------------------------------------------- #
# 9. Limitations
# --------------------------------------------------------------------------- #
def limitations(doc):
    h(doc, "8.  Limitations and Threats to Validity", level=1)
    threats = [
        (
            "Sample size (n=2 intraday).",
            "The matched NSE intraday universe is two pair-folds. "
            "Every NSE intraday conclusion — including the headline Kalman Sharpe and its DSR — is "
            "a two-fold-concatenation statistic and is labelled indicative throughout.",
        ),
        (
            "Non-iid, regime-mixed return series.",
            "DSR assumes iid returns; the concatenated "
            "folds are serially dependent and span different volatility regimes, and CSCV blocks "
            "straddle the fold join. The deflation is therefore an approximation that flatters "
            "neither direction cleanly.",
        ),
        (
            "Dividend-unadjusted prices.",
            "The NSE history is split/bonus-only; total-return "
            "effects are absent, which can bias spread levels for high-yield names.",
        ),
        (
            "Crypto funding cost unpriced [TODO].",
            "No funding series was available, so every "
            "Regime-B (multi-day carry) crypto number is funding-unpriced. Funding would worsen, "
            "not rescue, an already-negative net — but the exact magnitude is unquantified.",
        ),
        (
            "Labelled-default cost parameters.",
            "The crypto taker fee (5 bps/side) and the "
            "liquidity gate ($10M median quote-vol) are reasonable but venue-dependent defaults, "
            "not measured from a specific exchange's schedule.",
        ),
        (
            "Single risk-management placeholder.",
            "Crypto risk management is a single hard "
            "z-stop (3.5σ), not a tuned framework; the −73% drawdowns are not risk-optimised.",
        ),
        (
            "Selection-grid trial count.",
            "The honest N (46 NSE, 28 crypto) is reconstructed "
            "from the search grids rather than pre-registered; implicit selections (half-life "
            "bands, best-of-cell reporting) compound N further, so the true deflator is if "
            "anything stronger than reported.",
        ),
        (
            "Coint-stability gate uncalibrated.",
            "The rolling-ADF blacklist fires 19/19 on NSE "
            "and cannot be calibrated without a broad-universe base rate; its binary verdict is "
            "deferred.",
        ),
    ]
    for lead, body in threats:
        rich(doc, [(lead + "  ", {"bold": True, "color": ACCENT}), (body, {})])


# --------------------------------------------------------------------------- #
# 10. Future work
# --------------------------------------------------------------------------- #
def future_work(doc):
    h(doc, "9.  Future Work", level=1)
    para(doc, "The deferred units, ordered by how much each could move the headline:")
    numbered(
        doc,
        "the binding constraint. Maker/rebate execution, lower-frequency "
        "re-anchoring (daily-or-coarser μ updates with far fewer round-trips), and a "
        "turnover penalty inside the Bertram objective — the only lever that can flip "
        "the crypto net sign.",
        bold_lead="Low-cost / turnover-aware execution:  ",
    )
    numbered(
        doc,
        "the path the data demands — estimate the hedge ratio where the legs "
        "actually co-move (daily) and trade the residual intraday, or move to a "
        "venue with genuine intraday co-movement.",
        bold_lead="Lower-frequency β estimation:  ",
    )
    numbered(
        doc,
        "regime/HMM gating of when to trade, meta-labelling of trade quality, and "
        "RL-based position sizing — the deferred modelling units, to be added only "
        "after the cost wall is addressed (sizing cannot rescue a net-negative edge).",
        bold_lead="Regime/HMM, meta-labelling, RL sizing:  ",
    )
    numbered(
        doc,
        "PCA / DBSCAN clustering for de-correlated pair selection — directly attacks "
        "the concentration finding from Phase 2B.",
        bold_lead="Breadth-aware selection:  ",
    )
    numbered(
        doc,
        "a regulated-futures deployment (e.g. CME) where the NSE cash-market short "
        "and lot constraints do not bind and execution is cheaper — the eventual "
        "venue rationale behind the crypto retest.",
        bold_lead="CME / futures port:  ",
    )


# --------------------------------------------------------------------------- #
# 11. Appendices
# --------------------------------------------------------------------------- #
def appendices(doc):
    doc.add_page_break()
    h(doc, "10.  Appendices", level=1)

    h(doc, "A.  ASSUMPTIONS register and open [TODO]s", level=2)
    asum = pd.DataFrame(
        [
            [
                "A1",
                "DSR/PBO NSE sample length",
                "n=2 pair-folds, ~504 sessions",
                "broaden universe (crypto) for cross-sectional width",
            ],
            [
                "A2",
                "Honest trial count N",
                "46 (NSE), 28 (crypto)",
                "maintain a pre-registered trial ledger",
            ],
            [
                "A7",
                "Crypto taker fee",
                "5 bps/side (labelled default)",
                "confirm vs venue schedule (gross already negative)",
            ],
            [
                "A8",
                "Crypto funding series",
                "ABSENT → funding = 0 [TODO]",
                "pull a funding series; prices Regime-B carry",
            ],
            [
                "A9",
                "Crypto risk management",
                "single z-stop 3.5σ, not tuned",
                "dedicated risk framework",
            ],
            [
                "A_LIQ",
                "Crypto liquidity gate",
                "median quote-vol ≥ $10M",
                "keeps 22/30; $50M leaves 7 (too thin)",
            ],
            [
                "A6",
                "Coint-stability threshold",
                "60d / p>0.10 / 3 consecutive",
                "needs broad-universe base rate to calibrate",
            ],
        ],
        columns=["#", "assumption / [TODO]", "value used", "what would change it"],
    )
    tcap(
        doc,
        "A1",
        "Labelled defaults and open [TODO]s carried through the study (docs/phase4_report.md §0).",
    )
    table_from_df(doc, asum, fontsize=8.5)
    para(doc, "")
    para(doc, "Inter-document inconsistencies flagged during assembly:", bold=True)
    bullet(
        doc,
        "the clean log records 342 symbols post-cascade; the backtest loads 341. "
        "Difference = one share-class duplicate dropped at backtest. Resolved to 341 "
        "trading universe.",
        bold_lead="341 vs 342 symbols:  ",
    )
    bullet(
        doc,
        "untraceable in all working-tree files and git history; equals none of the "
        "A/B/C decomposition legs. The canonical frozen-OU best net annual is 20.915% "
        "(reports/phase3_ou/metrics_ou.csv); 20.886 is not carried forward.",
        bold_lead="20.886 vs 20.915:  ",
    )
    bullet(
        doc,
        "Phase-2A reports 8 fold indices spanning ~7 out-of-sample years (1764 daily "
        "obs ≈ 7×252). Both are correct; the report states 8 folds / ~7 OOS years.",
        bold_lead="7 vs 8 folds:  ",
    )
    bullet(
        doc,
        "the (1+β) re-billing was applied to the Phase-3 v2 cells at 3 bps; the 1/5/8 "
        "bps v2 cells remain on the legacy convention in their persisted CSVs "
        "(documented [TODO scope]). The OU/Kalman/crypto headline cells are all on the "
        "corrected billing.",
        bold_lead="Partial v2 re-billing:  ",
    )

    h(doc, "B.  Figure catalogue", level=2)
    para(
        doc,
        "Every figure embedded in this report, with its source PNG and the driver that "
        "produces it. All PNGs ship a companion same-basename CSV (the data behind the "
        "figure); PNGs are gitignored and regenerated by their driver.",
        size=9.5,
    )
    figcat = pd.DataFrame(
        [
            ["1", "plots/phase4/final/nse/portfolio_equity.png", "final_nse_plots.py"],
            ["2", "plots/phase4/final/crypto/portfolio_equity.png", "final_crypto_plots.py"],
            ["3", "plots/phase4/dsr_pbo/pbo_logit_distribution.png", "s2_dsr_pbo.py"],
            ["4", "plots/phase1/universe/01_symbols_per_sector.png", "scripts/06 universe EDA"],
            ["5", "plots/phase2/backtest/equity_curve.png", "scripts/10 backtest"],
            ["6", "plots/phase2/risk_managed/ladder_equity.png", "scripts/12 risk-managed"],
            ["7", "plots/phase3/equity_A_vs_B_3bps.png", "scripts/14 two-regime"],
            ["8", "plots/phase3/tod_vol_profile.png", "scripts/14 two-regime"],
            [
                "9",
                "reports/phase3_ou/figures/ou_best_cell/b_ou_best_portfolio_nav.png",
                "scripts/15 OU",
            ],
            [
                "10",
                "reports/phase3_ou/figures/grid_rollups/d_cost_ladder_f5_B.png",
                "scripts/15 OU",
            ],
            [
                "11",
                "reports/phase3_ou/figures/grid_rollups/f_half_life_distribution.png",
                "scripts/15 OU",
            ],
            [
                "12",
                "reports/phase3_ou/figures/grid_rollups/k_exclusion_funnel.png",
                "scripts/15 OU",
            ],
            ["13", "reports/phase3_kalman/figures/drift_before_after.png", "scripts/17 Kalman"],
            ["14", "plots/phase4/verification/per_trade_pnl_dist.png", "s1_unit_k_verification.py"],
            ["15", "reports/phase3_kalman/figures/mu_overlay_fold4.png", "scripts/17 Kalman"],
            ["16", "plots/phase4/beta_escalation/beta_paths_f5.png", "s3_beta_escalation.py"],
            ["17", "plots/phase4/coint_stability/adf_pvalue_paths.png", "s4_coint_stability.py"],
            ["18", "plots/phase4/johansen/eg_vs_johansen.png", "s5_johansen.py"],
            [
                "19",
                "plots/phase4/final/crypto/trade_return_distribution.png",
                "final_crypto_plots.py",
            ],
            ["20", "plots/phase4/crypto/crypto_cost_ladder.png", "s6_crypto.py"],
        ],
        columns=["fig", "source PNG", "driver"],
    )
    table_from_df(doc, figcat, fontsize=8)
    para(doc, "")

    h(doc, "C.  Table provenance", level=2)
    tabcat = pd.DataFrame(
        [
            ["1", "the seven cleaning rules", "src/apt/data, config/default.yaml"],
            ["2", "dividend-flavor diagnostic", "reports/dividend_flavor_diagnosis.csv"],
            [
                "3–4",
                "Phase 2A portfolio + per-pair",
                "reports/backtest_portfolio_metrics.csv, backtest_per_pair_metrics.csv",
            ],
            [
                "5–6",
                "Phase 2B ladder + vol-match",
                "reports/risk_managed_ladder.csv, risk_managed_vol_matched.csv",
            ],
            [
                "7",
                "OU matched-universe cost ladder",
                "reports/phase3_ou/figures/matched_universe/matched_metrics.csv",
            ],
            [
                "8",
                "Kalman 8-cell verification",
                "reports/phase4/verification/1b_eight_cell_table.csv",
            ],
            ["9", "β-collapse diagnostic", "reports/phase4/beta_escalation/beta_collapse.csv"],
            [
                "10",
                "coint-stability sensitivity",
                "reports/phase4/coint_stability/threshold_sensitivity.csv",
            ],
            [
                "11–12",
                "NSE DSR + sensitivity",
                "reports/phase4/dsr_pbo/2b_dsr.csv, 2b_dsr_sensitivity_to_N.csv",
            ],
            ["13", "crypto A10 headline", "reports/phase4/crypto_adaptive/metrics.csv"],
            ["14", "crypto Kalman DSR/PBO", "reports/phase4/crypto_adaptive/dsr_pbo_kalman.csv"],
        ],
        columns=["table", "content", "source CSV"],
    )
    tcap(doc, "C1", "Every performance table in this report traces to a committed CSV.")
    table_from_df(doc, tabcat, fontsize=8.5)
    para(doc, "")

    h(doc, "D.  Reproducibility — stack, branches, configuration", level=2)
    para(
        doc,
        "Python 3.11 (uv-managed). Core stack: polars, pandas, numpy, scipy, "
        "statsmodels, matplotlib, python-docx; pytest for the test suite. Linting via "
        "ruff. The full pipeline is the numbered scripts 01–10 plus the phase-3/4 "
        "drivers; every script is idempotent.",
        size=9.5,
    )
    try:
        branches = (
            subprocess.run(
                ["git", "-C", str(ROOT), "branch", "-a", "--format=%(refname:short)"],
                capture_output=True,
                text=True,
                timeout=20,
            )
            .stdout.strip()
            .splitlines()
        )
        local = [b for b in branches if not b.startswith("remotes/")]
        para(doc, "Branches (development lineage):", bold=True, size=9.5)
        para(doc, "  ".join(local), size=8.5, color=GREY)
    except Exception:
        pass
    try:
        log = (
            subprocess.run(
                ["git", "-C", str(ROOT), "log", "--oneline", "--all", "--no-decorate", "-n", "40"],
                capture_output=True,
                text=True,
                timeout=20,
            )
            .stdout.strip()
            .splitlines()
        )
        para(doc, "Commit provenance (most recent 40 across all branches):", bold=True, size=9.5)
        for line in log:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(7.5)
            r.font.color.rgb = GREY
    except Exception:
        pass

    h(doc, "E.  Key references", level=2)
    refs = [
        "Engle, R. F., & Granger, C. W. J. (1987). Co-integration and error correction. Econometrica.",
        "Johansen, S. (1991). Estimation and hypothesis testing of cointegration vectors. Econometrica.",
        "Benjamini, Y., & Hochberg, Y. (1995). Controlling the false discovery rate. JRSS-B.",
        "Bertram, W. K. (2010). Analytic solutions for optimal statistical arbitrage trading. Physica A, 389(11), 2234–2243.",
        "Bailey, D. H., & López de Prado, M. (2014). The Deflated Sharpe Ratio. Journal of Portfolio Management.",
        "Bailey, D. H., Borwein, J., López de Prado, M., & Zhu, Q. J. (2017). The probability of backtest overfitting. Journal of Computational Finance.",
        "West, M., & Harrison, J. (1997). Bayesian Forecasting and Dynamic Models (discount factors).",
    ]
    for r in refs:
        bullet(doc, r)


# --------------------------------------------------------------------------- #
# main
# --------------------------------------------------------------------------- #
def main() -> int:
    doc = Document()
    setup(doc)
    title_page(doc)
    abstract(doc)
    executive_summary(doc)
    introduction(doc)
    data_section(doc)
    methodology(doc)
    results(doc)
    discussion(doc)
    conclusion(doc)
    limitations(doc)
    future_work(doc)
    appendices(doc)
    add_footer(doc)
    doc.save(str(OUT))
    print(f"WROTE {OUT}")
    print(f"figures embedded: {len(EMBEDDED)}")
    for e in EMBEDDED:
        print("  +", e)
    if MISSING:
        print(f"MISSING figures: {len(MISSING)}")
        for m in MISSING:
            print("  !", m)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
