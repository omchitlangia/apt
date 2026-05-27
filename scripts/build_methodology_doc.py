"""One-shot builder for the Checkpoint 1 methodology Word doc."""

from __future__ import annotations

import contextlib
import csv
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "reports" / "APT_Methodology_Checkpoint_1.docx"


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_image(doc: Document, rel_path: str, caption: str, width_in: float = 5.5) -> None:
    path = ROOT / rel_path
    if not path.exists():
        doc.add_paragraph(f"[MISSING IMAGE: {rel_path}]")
        return
    doc.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def csv_to_table(
    doc: Document,
    csv_path: Path,
    max_rows: int | None = None,
    columns: list[str] | None = None,
    col_renames: dict[str, str] | None = None,
    numeric_fmt: dict[str, str] | None = None,
) -> None:
    with csv_path.open() as f:
        reader = csv.DictReader(f)
        all_rows = list(reader)
        header = reader.fieldnames or []
    if columns:
        header = [c for c in columns if c in header]
    if max_rows is not None:
        all_rows = all_rows[:max_rows]
    display_header = [(col_renames or {}).get(h, h) for h in header]

    tbl = doc.add_table(rows=1 + len(all_rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(display_header):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(all_rows, start=1):
        for j, h in enumerate(header):
            val = row.get(h, "")
            fmt = (numeric_fmt or {}).get(h)
            if fmt and val not in ("", None):
                with contextlib.suppress(TypeError, ValueError):
                    val = fmt.format(float(val))
            tbl.rows[i].cells[j].text = str(val)


def main() -> None:
    doc = Document()

    # Title (plain — no cover page, no styled theme)
    title = doc.add_paragraph()
    run = title.add_run("APT Methodology — Checkpoint 1 (Draft)")
    run.bold = True
    run.font.size = Pt(16)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "First-pass draft for manual revision. Numbers pulled from Phase 1 / Phase 2A / "
        "Phase 2B report artifacts in this repo; [TODO] markers flag items needing review."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)

    # ---------------- 1. Data and adjustments ----------------
    doc.add_heading("1. Data and adjustments", level=1)
    doc.add_paragraph(
        "Daily OHLCV CSVs for the NSE Nifty-500 universe (read-only at "
        "/Data6/db/) entered the pipeline with 492 distinct symbols and "
        "1,616,202 raw price rows. After back-adjustment for corporate actions "
        "and seven cleaning rules (calendar filter; residual split repair; "
        "phantom-history trimming at a 0.65 jump threshold; structural-event "
        "handling; 60-day rolling-median ADV liquidity floor of ₹1 crore; a "
        "10-day contiguity gap cap; and a 756-day minimum history requirement), "
        "the cleaned panel covers 342 symbols and 794,220 daily rows over "
        "2004-06-23 to 2021-06-23 (~17 years; weekday-trading-day basis ~252 "
        "rows/year)."
    )
    doc.add_paragraph(
        "We use SPLIT- and BONUS-adjusted close prices only. We do NOT apply "
        "dividend or total-return adjustment. The reason is mechanical: a "
        "total-return series adds back a path-dependent dividend stream whose "
        "shape differs across names, so the spread between two cointegrated "
        "tickers picks up a synthetic drift driven by their differing payout "
        "policies rather than by any genuine economic dislocation. That drift "
        "corrupts (a) the Engle–Granger cointegration test, (b) the estimated "
        "hedge ratio β, and (c) the rolling z-score bands the signal rests on. "
        "Using split-only prices keeps the spread a clean function of the two "
        "traded instruments."
    )
    doc.add_paragraph(
        "We verified the adjustment flavor of the raw vendor feed with a "
        "ratio-stability diagnostic against yfinance: comparing the raw close "
        "to yfinance's split-only Close and yfinance's total-return Adj Close "
        "on five high-dividend names (ITC, COALINDIA, HINDPETRO, NTPC, "
        "POWERGRID). For all five, the coefficient of variation of "
        "raw/yfinance-Close was at least an order of magnitude smaller than "
        "raw/yfinance-Adj-Close — verdict 'split_only' with high confidence "
        "on four and medium on NTPC. The diagnostic output is preserved at "
        "reports/dividend_flavor_diagnosis.csv."
    )
    doc.add_paragraph(
        "The consequence we accept and disclose: dividend cash flows on held "
        "positions are not modeled. For a market-neutral pair the long leg "
        "receives and the short leg pays on ex-date, so the bias is "
        "near-zero on average over a balanced book, but for any individual "
        "pair the residual drag/credit can run ~1–2 %/yr. We treat this as a "
        "stated small bias rather than capitalizing it into P&L."
    )

    add_image(
        doc,
        "plots/phase1/universe/02_history_length_distribution.png",
        "Figure 1.1 — History-length distribution across the cleaned universe.",
    )
    add_image(
        doc,
        "plots/phase1/universe/03_coverage_heatmap.png",
        "Figure 1.2 — Symbol × time coverage heatmap (post-cleaning).",
    )

    # ---------------- 2. Transaction cost economics ----------------
    doc.add_heading("2. Transaction cost economics", level=1)
    doc.add_paragraph(
        "Backtest cost model: a static round-trip charge of 25 bps per leg "
        "(config/default.yaml → backtest.cost_bps_per_leg). A pair trade "
        "incurs the charge on both legs at entry and at exit, i.e. 4 × 25 = "
        "100 bps of round-trip frictional cost per pair round-trip in "
        "log-cost space. The 25 bps figure sits in the middle of a realistic "
        "Indian-cash-equity range of roughly 15–30 bps, decomposing into "
        "brokerage, STT (securities transaction tax), exchange/clearing fees, "
        "and a slippage allowance for the bid-ask plus realistic fill quality "
        "at the mid-cap end of the universe."
    )
    doc.add_paragraph(
        "All P&L numbers in this report are reported both gross and net of "
        "cost; the headline performance figures in Section 5 and the entire "
        "Section 7 risk-management discussion are net. The bar a pair has to "
        "clear is therefore the gross statistical edge minus this frictional "
        "drag — at 100 bps per round-trip and the observed ~14–16 day mean "
        "holding period, only edges of meaningful magnitude survive into net "
        "P&L. The cost model is acknowledged to be static; real slippage "
        "spikes under stress (gap opens, illiquid names), which is one of the "
        "items the Phase 2B risk overlay was designed to probe."
    )

    # ---------------- 3. Portfolio weight logic across pairs ----------------
    doc.add_heading("3. Portfolio weight logic across pairs", level=1)
    doc.add_paragraph(
        "Phase 2A sizing is deliberately the simplest possible: fixed equal "
        "notional per active pair. When a pair has an open position the "
        "portfolio commits one unit of notional to it; idle pairs sit at zero. "
        "The portfolio daily return is the sum of the per-pair daily log "
        "returns net of cost, with no leverage scaling, no vol targeting, and "
        "no inter-pair re-weighting. Idle capital earns 0 % — there is no "
        "cash-sweep return modeled."
    )
    doc.add_paragraph(
        "Across the 7-year walk-forward, the portfolio is in at least one "
        "position roughly one third of the available time. That low "
        "deployment ratio is the dominant reason the absolute return is "
        "modest in relation to the per-trade edge: the engine is selective, "
        "not always-on. It also means raw portfolio vol understates the vol "
        "during deployed windows, which is the lens Phase 2B revisits."
    )
    doc.add_paragraph(
        "Caveat: pair P&L streams are not independent. The pair universe is "
        "financials-heavy and shares legs across pairs (HDFCBANK alone touches "
        "roughly a third of distinct tradeable pairs across the backtest, per "
        "reports/backtest_caveats.txt). Treating the portfolio as a naive sum "
        "of independent pair P&L's overstates the diversification benefit and "
        "inflates Sharpe; Section 7 walks through how this concentration "
        "drove the 2018-19 loss cluster."
    )

    # ---------------- 4. Intra-pair weighting (the two legs) ----------------
    doc.add_heading("4. Intra-pair weighting (the two legs)", level=1)
    doc.add_paragraph(
        "Within a pair the two legs are weighted by the hedge ratio β from "
        "an Engle–Granger OLS regression on log prices:"
    )
    p = doc.add_paragraph()
    p.add_run("    spread_t = log(P1_t) − β · log(P2_t) − α").italic = True
    doc.add_paragraph(
        "We run the OLS in both directions (P1 on P2 and P2 on P1), test both "
        "residual series with ADF, and select the direction whose residual "
        "rejects the unit root more strongly. The β (and intercept α) we end "
        "up with are the leg weights: a one-unit notional position in the "
        "leading leg is paired with a β-unit notional position of opposite "
        "sign in the second leg."
    )
    doc.add_paragraph(
        "Critically, (α, β) are estimated on the TRAINING fold only and FROZEN "
        "for the entire 252-day test fold. They are never re-fit on test data, "
        "and the rolling z-score that drives entries and exits uses the "
        "frozen β through the test window. This is the pair-level "
        "leakage guard: every parameter that touches a trade decision was "
        "determined before that trade's price history was seen. Re-selection "
        "happens at the fold boundary every 252 days, at which point a new "
        "training window can produce a different β or drop the pair entirely "
        "from the active set."
    )

    # ---------------- 5. Per-pair trade reporting + aggregate performance ----------------
    doc.add_heading("5. Per-pair trade reporting + aggregate performance", level=1)
    doc.add_paragraph(
        "Phase 2A ran a walk-forward backtest across 8 annual folds (test "
        "window = 252 trading days, training window = 1008 days, fold step = "
        "252 days). The selection pipeline produced 30 pair-fold units across "
        "the 8 folds, generating 198 trades (147 mean-revert exits, 34 stop "
        "exits at |z|>3.5, 11 fold-boundary force-closes, 6 time-stop exits)."
    )
    doc.add_paragraph(
        "Headline portfolio metrics (reports/backtest_portfolio_metrics.csv, net of 25 bps/leg):"
    )
    csv_to_table(
        doc,
        ROOT / "reports/backtest_portfolio_metrics.csv",
        numeric_fmt={
            "gross": "{:.3f}",
            "net": "{:.3f}",
        },
    )
    add_caption(doc, "Table 5.1 — Walk-forward portfolio metrics, gross vs net.")

    doc.add_paragraph(
        "Net Sharpe is 1.08 and net annualised return is 17.76 % over the "
        "7-year out-of-sample window, with realised net annualised vol of "
        "15.18 % and a peak-to-trough net drawdown of -17.45 %. Per-trade "
        "win rate (net basis) is 66.2 %, and 74.2 % of all exits are "
        "mean-reversion exits — i.e. the trading thesis (mean reversion to "
        "the cointegration relationship) is what closes the median trade, "
        "rather than a stop-out or a fold-boundary force-close."
    )
    doc.add_paragraph(
        "Per-pair, the tradeable set splits cleanly into earners and "
        "underperformers. Top earners by net annualised return:"
    )
    csv_to_table(
        doc,
        ROOT / "reports/phase2_per_pair.csv",
        max_rows=8,
        columns=[
            "pair",
            "sector",
            "n_trades",
            "n_obs",
            "net_total_pct",
            "net_ann_pct",
            "net_sharpe",
            "net_max_drawdown_pct",
            "win_rate_net",
            "avg_holding_days",
            "n_exits_mean_revert",
            "n_exits_stop",
        ],
        col_renames={
            "net_total_pct": "net_total_%",
            "net_ann_pct": "net_ann_%",
            "net_max_drawdown_pct": "net_maxDD_%",
            "win_rate_net": "win_rate",
            "avg_holding_days": "avg_hold_d",
            "n_exits_mean_revert": "exits_MR",
            "n_exits_stop": "exits_stop",
        },
        numeric_fmt={
            "net_total_pct": "{:.1f}",
            "net_ann_pct": "{:.1f}",
            "net_sharpe": "{:.2f}",
            "net_max_drawdown_pct": "{:.1f}",
            "win_rate_net": "{:.2f}",
            "avg_holding_days": "{:.1f}",
        },
    )
    add_caption(
        doc,
        "Table 5.2 — Top per-pair trade summary (full table in reports/phase2_per_pair.csv).",
    )

    add_image(
        doc,
        "plots/phase2/backtest/equity_curve.png",
        "Figure 5.1 — Phase 2A walk-forward portfolio equity curve (gross and net of 25 bps/leg).",
    )
    add_image(
        doc,
        "plots/phase2/pairs/per_pair_overview.png",
        "Figure 5.2 — Per-pair contribution grid (Phase 2A diagnostic).",
    )

    # ---------------- 6. Cointegration and its breakdown ----------------
    doc.add_heading("6. Cointegration and its breakdown", level=1)
    doc.add_paragraph(
        "Selection runs once per fold on the 1008-day training window. The "
        "machinery: (a) restrict candidate pairs to same-sector clusters from "
        "the ind_nifty500list mapping; (b) apply a loose correlation "
        "pre-filter on log-returns at ρ ≥ 0.50 over the prior 504 days, "
        "deliberately loose so that the cointegration stage has room to make "
        "multiple-testing inflation visible; (c) run Engle–Granger OLS in "
        "both directions and test the residual with ADF at p < 0.05; (d) "
        "apply a Benjamini–Hochberg FDR correction at α = 0.05 across the "
        "fold's candidate set; (e) apply robustness gates — half-life in "
        "[5, 60] trading days, Hurst exponent < 0.5, and cross-window "
        "stability between adjacent training folds."
    )
    doc.add_paragraph(
        "Funnel (full cointegration parquet, "
        "data/pairs/cointegrated_pairs.parquet): 234 candidate pairs passed "
        "the correlation pre-filter and entered cointegration testing; 57 "
        "rejected the unit root at raw p<0.05; only 15 survived the "
        "Benjamini–Hochberg FDR cut at α=0.05 — a clear illustration of "
        "the multiple-comparisons inflation the FDR step exists to control. "
        "After the half-life/Hurst/stability robustness gates, 7 pairs "
        "appeared in the final ranked report (reports/"
        "cointegrated_pairs_ranked.csv) for the most recent training window."
    )
    doc.add_paragraph(
        "Breakdown handling. Cointegration is a statistical relationship that "
        "decays. Three guards: per-fold re-selection every 252 days (a pair "
        "that fails the gates on the new training window simply does not "
        "trade in the new test fold); a z-score stop at |z|>3.5 that closes "
        "any trade where the spread has walked far enough from its training "
        "mean to suggest the relationship has broken; and a max-holding "
        "time-stop at min(60d, 3 × half_life), which prevents capital from "
        "sitting in a position that is simply not mean-reverting on the "
        "expected timescale."
    )
    add_image(
        doc,
        "plots/phase1/pairs/spread_ONGC_OIL.png",
        "Figure 6.1 — Representative spread & rolling z-score diagnostic (ONGC / OIL).",
    )
    add_image(
        doc,
        "plots/phase1/pairs/signal_ONGC_OIL.png",
        "Figure 6.2 — Entry/exit signal overlay for the same pair (Day 7 diagnostic).",
    )

    # ---------------- 7. Drawbacks, the loss cluster, and the failed risk layer ----------------
    doc.add_heading("7. Drawbacks, the loss cluster, and the failed risk layer", level=1)
    doc.add_paragraph(
        "The honest part of the result lives in the 2018-19 fold. Under the "
        "Phase 2A equal-notional sizing, a cluster of correlated PSU-bank "
        "pairs (BANKINDIA/IDBI, IDBI/CANBK, IDBI/BANKINDIA, IDBI/UNIONBANK, "
        "UNIONBANK/IDBI, BANKINDIA/SBIN, BANKINDIA/IOB — see "
        "reports/phase2_per_pair.csv) broke down together as PSU-bank credit "
        "stress and the IL&FS shock pulled correlated leg pairs apart "
        "simultaneously. Because the equal-notional book treated these as "
        "independent pairs, the fold absorbed the full sum of those concurrent "
        "drawdowns; the fold-level net P&L was deeply negative "
        "[TODO: confirm precise 2018-19 fold net % from "
        "reports/backtest_portfolio_daily.csv — user-stated figure was ~-17 %]."
    )
    doc.add_paragraph(
        "Phase 2B was the structured attempt to fix this with a risk overlay: "
        "fixed-fractional sizing (R1), per-pair caps (R2), sector/cluster "
        "caps (R3), and a relationship-driven pair-kill switch (R4). The "
        "ablation ladder (reports/risk_managed_ladder.csv) is in Table 7.1."
    )
    csv_to_table(
        doc,
        ROOT / "reports/risk_managed_ladder.csv",
        columns=[
            "label",
            "kill_mode",
            "cluster_cap",
            "n_trades",
            "net_ann_pct",
            "ann_vol_pct",
            "net_sharpe",
            "max_drawdown_pct",
            "fold_2018_19_net_pct",
            "carrier_PFC_SBIN_net_pct",
            "carrier_ONGC_OIL_net_pct",
        ],
        col_renames={
            "net_ann_pct": "net_ann_%",
            "ann_vol_pct": "ann_vol_%",
            "max_drawdown_pct": "maxDD_%",
            "fold_2018_19_net_pct": "2018-19_net_%",
            "carrier_PFC_SBIN_net_pct": "PFC/SBIN_net_%",
            "carrier_ONGC_OIL_net_pct": "ONGC/OIL_net_%",
        },
        numeric_fmt={
            "net_ann_pct": "{:.2f}",
            "ann_vol_pct": "{:.2f}",
            "net_sharpe": "{:.3f}",
            "max_drawdown_pct": "{:.2f}",
            "fold_2018_19_net_pct": "{:.2f}",
            "carrier_PFC_SBIN_net_pct": "{:.2f}",
            "carrier_ONGC_OIL_net_pct": "{:.2f}",
            "cluster_cap": "{:.2f}",
        },
    )
    add_caption(
        doc,
        "Table 7.1 — Risk-overlay ablation ladder (R0 = Phase 2A baseline; "
        "R3 = cluster-capped; R4 = + relationship pair-kill).",
    )

    doc.add_paragraph(
        "Read naively, R3 looks like a win on the 2018-19 metric (+1.32 % "
        "net vs R0's -3.47 % net on the fold) and on max-DD (-4.24 % vs "
        "-17.45 %). But the comparison is unfair: R3 is dramatically "
        "de-risked (raw ann vol 1.47 % vs R0's 15.18 %), so a same-Sharpe "
        "comparison cannot use raw curves. We re-ran R0 and R3 through a "
        "causal vol-target overlay at 10 % annualised vol with a 3× leverage "
        "cap (apt.backtest.apply_vol_target_overlay; "
        "reports/risk_managed_vol_matched.csv):"
    )
    csv_to_table(
        doc,
        ROOT / "reports/risk_managed_vol_matched.csv",
        columns=[
            "rung",
            "raw_ann_vol_pct",
            "vol_targeted_ann_vol_pct",
            "vol_targeted_ann_return_pct",
            "vol_targeted_sharpe",
            "vol_targeted_max_dd_pct",
            "vol_targeted_fold_2018_19_pct",
            "mean_leverage",
            "max_leverage_used",
        ],
        col_renames={
            "raw_ann_vol_pct": "raw_vol_%",
            "vol_targeted_ann_vol_pct": "vt_vol_%",
            "vol_targeted_ann_return_pct": "vt_ret_%",
            "vol_targeted_sharpe": "vt_Sharpe",
            "vol_targeted_max_dd_pct": "vt_maxDD_%",
            "vol_targeted_fold_2018_19_pct": "vt_18-19_%",
            "mean_leverage": "mean_lev",
            "max_leverage_used": "max_lev",
        },
        numeric_fmt={
            "raw_ann_vol_pct": "{:.2f}",
            "vol_targeted_ann_vol_pct": "{:.2f}",
            "vol_targeted_ann_return_pct": "{:.2f}",
            "vol_targeted_sharpe": "{:.3f}",
            "vol_targeted_max_dd_pct": "{:.2f}",
            "vol_targeted_fold_2018_19_pct": "{:.2f}",
            "mean_leverage": "{:.2f}",
            "max_leverage_used": "{:.2f}",
        },
    )
    add_caption(
        doc,
        "Table 7.2 — Vol-matched comparison at 10 % target / 3× leverage cap.",
    )
    add_image(
        doc,
        "plots/phase2/risk_managed/final_equity_curve.png",
        "Figure 7.1 — Final vol-matched equity curves: R0 vs R3 at the same realised vol.",
    )
    doc.add_paragraph(
        "The decisive finding: at the 3× leverage cap, R0 and R3 Sharpes are "
        "essentially tied (0.941 vs 0.930). R3 is leverage-pinned (mean "
        "leverage 2.98, hard against the 3× cap), and the cluster cap that "
        "looked like an edge on raw curves was a deployment artifact — a "
        "smaller book mechanically takes a smaller dollar loss in the bad "
        "fold. Once you normalise for the vol you can actually deploy, the "
        "cluster cap has no signal. Worse, R3 guts the carriers: PFC/SBIN's "
        "net contribution collapses from +90.6 % to +4.6 %, and ONGC/OIL "
        "from +21.6 % to +2.4 %. The cap throws away the trades that paid "
        "for the strategy."
    )
    doc.add_paragraph(
        "What R3 does keep is genuine tail protection on the bad fold "
        "(2018-19 vol-matched: +4.01 % vs R0 -6.68 %) and a tighter "
        "max-drawdown (-12.19 % vs -18.57 %). So the cluster cap is a real "
        "tail-risk knob — it just cannot manufacture additional risk-adjusted "
        "return on this universe."
    )
    add_image(
        doc,
        "plots/phase2/risk_managed/drawdown_per_rung.png",
        "Figure 7.2 — Drawdown trajectories across the R0→R4 ablation ladder.",
    )
    add_image(
        doc,
        "plots/phase2/risk_managed/cluster_exposure_R4.png",
        "Figure 7.3 — Per-cluster (sector) exposure timeline under R4.",
    )
    add_image(
        doc,
        "plots/phase2/risk_managed/carriers_R0_vs_R4.png",
        "Figure 7.4 — Per-trade P&L of the carrier pairs (PFC/SBIN, ONGC/OIL) under R0 vs R4.",
    )
    add_image(
        doc,
        "plots/phase2/risk_managed/cluster_cap_sweep.png",
        "Figure 7.5 — Cluster-cap sweep on R4 (none → 5 % → 3 %).",
    )
    doc.add_paragraph(
        "Conclusion to state plainly. Blunt risk caps cannot manufacture "
        "robustness on a thin, concentrated universe. The selection + "
        "cointegration + signal framework is sound, but the equities surface "
        "we have is too narrow: too few independent pairs, too many shared "
        "legs, and — critically — no live multi-day shortability on Indian "
        "cash equity, which makes the entire Phase 2A P&L an upper-bound "
        "POC rather than a deployable strategy. The natural next venues are "
        "crypto and US futures: more breadth, genuine shortability, and "
        "enough deployable raw vol that a vol-target overlay does not collide "
        "with the leverage cap."
    )
    doc.add_paragraph("Honest caveat list (carried forward into the next phase):")
    bullets = [
        "Multi-day shorting on Indian cash equity is not executable — this is "
        "an upper-bound POC P&L, not a live-tradeable strategy.",
        "Heavy shared-leg / financials concentration: pair returns are not "
        "independent and Sharpe is overstated when treated as if they were.",
        "Thin per-fold pair count (typically 0–7 tradeable pairs) → wide "
        "error bars on every metric.",
        "Dividends not modeled (split-only prices); ~1–2 %/yr per-pair "
        "residual bias on individual names.",
        "Cost model is a static per-leg bps figure; real slippage spikes "
        "under stress and the 25 bps default sits mid-range.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
